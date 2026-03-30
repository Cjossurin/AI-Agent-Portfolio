"""
Document writer utilities — generates DOCX and PDF dispute letters.

Uses python-docx for Word document creation and docx2pdf for
Windows-native PDF conversion via Word COM automation.
"""

from __future__ import annotations

import logging
from pathlib import Path

import re

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models.schemas import PersonalInfo

logger = logging.getLogger(__name__)


def write_docx(
    content: str,
    output_path: str | Path,
    personal_info: PersonalInfo,
    bureau: str,
    current_date: str,
) -> Path:
    """Create a formatted Word document for a dispute letter.

    Args:
        content: The full body text of the dispute letter.
        output_path: Destination file path for the .docx file.
        personal_info: Consumer's personal information for the header block.
        bureau: Target credit bureau name.
        current_date: Human-readable date string (e.g. 'February 28, 2026').

    Returns:
        Path to the created .docx file.
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── Sender block (consumer info) ─────────────────────────────────────
    sender_lines = [
        personal_info.full_name,
        personal_info.current_address,
    ]
    if personal_info.phone:
        sender_lines.append(personal_info.phone)
    if personal_info.email:
        sender_lines.append(personal_info.email)
    if personal_info.ssn_last4:
        sender_lines.append(f"SSN (last 4): ***-**-{personal_info.ssn_last4}")
    if personal_info.date_of_birth:
        sender_lines.append(f"DOB: {personal_info.date_of_birth}")

    for line in sender_lines:
        p = doc.add_paragraph(line)
        p.style.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    # ── Date ─────────────────────────────────────────────────────────────
    doc.add_paragraph("")  # spacer
    date_para = doc.add_paragraph(current_date)
    date_para.style.font.size = Pt(11)

    # ── Title ────────────────────────────────────────────────────────────
    title = doc.add_paragraph(f"{bureau} \u2013 Credit Report Dispute")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0] if title.runs else title.add_run()
    title_run.bold = True
    title_run.font.size = Pt(14)

    doc.add_paragraph("")  # spacer

    # ── Body paragraphs ──────────────────────────────────────────────────
    # Split on double newlines (paragraph breaks) so each logical block
    # becomes one Word paragraph with keep_together=True, preventing
    # any paragraph from being split across a page boundary.
    # Short ALL-CAPS blocks (only letters/spaces/&, 2–10 words, ≤ 80 chars)
    # are rendered as bold 12pt section headings with extra top spacing.
    # Blocks containing bullet lines (•) are rendered as individual tight-
    # spaced paragraphs so each bullet stays on its own line.
    for block in content.split("\n\n"):
        lines = [line.strip() for line in block.split("\n") if line.strip()]
        if not lines:
            continue
        clean = " ".join(lines)

        is_heading = bool(
            re.match(r'^[A-Z][A-Z &]+$', clean)
            and 2 <= len(clean.split()) <= 10
            and len(clean) <= 80
        )
        has_bullets = any(line.startswith("\u2022") for line in lines)

        if is_heading:
            p = doc.add_paragraph()
            run = p.add_run(clean)
            run.bold = True
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.keep_together = True
        elif has_bullets:
            for line in lines:
                p = doc.add_paragraph(line)
                p.style.font.size = Pt(11)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.keep_together = True
        else:
            p = doc.add_paragraph(clean)
            p.style.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.space_before = Pt(4)
            # Prevent the paragraph from being broken across a page.
            p.paragraph_format.keep_together = True

    # ── Signature block (dotted-line handwriting areas) ────────────────────
    doc.add_paragraph("")  # spacer
    closer = doc.add_paragraph("Sincerely,")
    closer.style.font.size = Pt(11)
    closer.paragraph_format.space_after = Pt(2)

    # Three blank lines for handwritten signature space
    for _ in range(3):
        spacer = doc.add_paragraph("")
        spacer.paragraph_format.space_after = Pt(0)
        spacer.paragraph_format.space_before = Pt(0)

    # Signature dotted line
    sig_line = doc.add_paragraph("Signature: " + "_" * 50)
    sig_line.style.font.size = Pt(11)
    sig_line.paragraph_format.space_after = Pt(6)

    # Printed name
    name_line = doc.add_paragraph(f"Printed Name: {personal_info.full_name}")
    name_line.style.font.size = Pt(11)
    name_line.paragraph_format.space_after = Pt(14)

    # Date dotted line
    date_sig = doc.add_paragraph("Date: " + "_" * 50)
    date_sig.style.font.size = Pt(11)
    date_sig.paragraph_format.space_after = Pt(6)

    # ── Save ─────────────────────────────────────────────────────────────
    doc.save(str(output_path))
    logger.info("Saved DOCX: %s", output_path)
    return output_path


def convert_to_pdf(docx_path: str | Path) -> Path:
    """Convert a .docx file to PDF using docx2pdf (Windows Word COM).

    Args:
        docx_path: Path to the source .docx file.

    Returns:
        Path to the generated .pdf file (same directory, same stem).

    Raises:
        RuntimeError: If conversion fails (e.g. Word is not installed).
    """
    from docx2pdf import convert

    docx_path = Path(docx_path)
    pdf_path = docx_path.with_suffix(".pdf")

    try:
        convert(str(docx_path), str(pdf_path))
        logger.info("Converted to PDF: %s", pdf_path)
    except Exception as exc:
        logger.error("PDF conversion failed for %s: %s", docx_path, exc)
        raise RuntimeError(
            f"Failed to convert {docx_path.name} to PDF. "
            f"Ensure Microsoft Word is installed. Error: {exc}"
        ) from exc

    return pdf_path
